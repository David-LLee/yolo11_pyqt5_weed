"""
基于YOLOv11-PyQt 的杂草识别、定位及状况监测系统程序
作者: David Lee
版本: v1.0
功能: 支持图像/相机检测，实时界面显示
"""

from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QMenu, QAction
from main_win.win6 import Ui_mainWindow
from PyQt5.QtCore import Qt, QPoint, QTimer, QThread, pyqtSignal
from PyQt5.QtGui import QImage, QPixmap, QIcon
import sys
import os
import json
import numpy as np
import torch
import time
import csv
import cv2
from datetime import datetime
from fpdf import FPDF       # 用于生成PDF报告
import traceback
import matplotlib.font_manager as fm

# YOLOv11依赖
from ultralytics import YOLO
from ultralytics.utils.plotting import Annotator, colors
from ultralytics.utils.torch_utils import select_device
from ultralytics.utils.capnums import Camera
from ultralytics.utils.CustomMessageBox import MessageBox


class DetThread(QThread):
    send_img = pyqtSignal(np.ndarray)  # 处理后图像
    send_raw = pyqtSignal(np.ndarray)  # 原始图像
    send_statistic = pyqtSignal(dict, list)
    send_msg = pyqtSignal(str)
    send_percent = pyqtSignal(int)
    send_fps = pyqtSignal(str)

    def __init__(self):
        super(DetThread, self).__init__()
        self.weights = './yolo11s.pt'
        self.current_weight = self.weights
        self.source = '0'
        self.conf_thres = 0.45
        self.iou_thres = 0.70
        self.jump_out = False
        self.is_continue = True
        self.percent_length = 1000
        self.rate_check = True
        self.rate = 100
        self.model = None
        self.device = None
        self.total_frames = 0

    @torch.no_grad()
    def run(self):
        try:
            self.device = select_device('')
            self.model = YOLO(self.weights)
            self.model.to(self.device)
            self.model.fuse()

            # 设置模型参数
            self.model.conf = self.conf_thres
            self.model.iou = self.iou_thres

            # 初始化数据源
            is_webcam = self.source.isnumeric() or self.source.lower().startswith(
                ('rtsp://', 'rtmp://', 'http://', 'https://'))

            # 强制指定摄像头后端（仅当source为数字时）
            if self.source.isnumeric():
                cap = cv2.VideoCapture(int(self.source), cv2.CAP_MSMF)  # 使用MSMF后端
                if not cap.isOpened():
                    self.send_msg.emit(f'错误: 无法打开摄像头 {self.source}')
                    return
                cap.release()
            else:
                cap = cv2.VideoCapture(self.source)
                if cap.isOpened():
                    self.total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                    cap.release()

            # 创建数据流（此处由YOLO管理摄像头）
            dataset = self.model.predict(
                source=self.source,
                stream=True,
                imgsz=640,
                show=False,
                verbose=False
            )

            count = 0
            start_time = time.time()

            while True:
                # 实时更新模型参数
                self.model.conf = self.conf_thres
                self.model.iou = self.iou_thres

                if self.jump_out:
                    break

                if self.current_weight != self.weights:
                    self.model = YOLO(self.weights)
                    self.model.to(self.device)
                    # 热更新后同步参数
                    self.model.conf = self.conf_thres  # 新增
                    self.model.iou = self.iou_thres  # 新增
                    self.current_weight = self.weights

                if self.is_continue:
                    try:
                        results = next(dataset)
                    except StopIteration:
                        self.send_percent.emit(0)
                        self.send_msg.emit('检测结束')
                        break

                    # 处理图像
                    raw_img = results.orig_img.copy()
                    processed_img = raw_img.copy()

                    # 统计和标注
                    statistic_dic = {}
                    annotator = Annotator(raw_img.copy(), line_width=3)  # 改为基于原始图像创建annotator

                    weed_coords = []  # 每次循环初始化坐标列表
                    for box in results.boxes:
                        # 过滤置信度低于阈值的检测框
                        if float(box.conf) < self.conf_thres:
                            continue
                        cls = int(box.cls)
                        conf = float(box.conf)
                        label = f"{self.model.names[cls]} {conf:.2f}"
                        statistic_dic[self.model.names[cls]] = statistic_dic.get(self.model.names[cls], 0) + 1

                        # 先绘制检测框
                        annotator.box_label(box.xyxy[0], label, color=colors(cls, True))

                        weed_classes = ["broadleaf_weed", "grass_weed", "sedge_weed"]
                        current_class = self.model.names[cls].lower()  # 统一转为小写

                        # 收集坐标（包含所有杂草类别）
                        if current_class in weed_classes:
                            x1, y1, x2, y2 = box.xyxy[0].cpu().numpy()
                            center_x = (x1 + x2) / 2
                            center_y = (y1 + y2) / 2
                            weed_coords.append((center_x, center_y))

                        # 仅针对weed绘制关键点
                        if self.model.names[cls].lower() == "Weed" and hasattr(box, 'keypoints'):  # 忽略大小写
                            keypoints = box.keypoints.xy[0].cpu().numpy()
                            h, w = annotator.im.shape[:2]
                            for kp in keypoints:
                                x = np.clip(int(kp[0]), 0, w - 1)  # 限制坐标范围
                                y = np.clip(int(kp[1]), 0, h - 1)
                                cv2.circle(annotator.im, (x, y), 8, (0, 0, 255), -1)

                    # 发送信号时添加坐标参数
                    self.send_statistic.emit(statistic_dic, weed_coords)  # 修改emit参数

                    processed_img = annotator.result()  # 最后获取结果


                    # 发送信号
                    self.send_img.emit(processed_img)
                    self.send_raw.emit(raw_img)

                    # 计算帧率
                    count += 1
                    if count % 30 == 0:
                        fps = int(30 / (time.time() - start_time))
                        self.send_fps.emit(f'FPS: {fps}')
                        start_time = time.time()

                    # 更新进度条（仅限视频文件）
                    if self.total_frames > 0:
                        percent = int(count / self.total_frames * self.percent_length)
                        self.send_percent.emit(percent)

                    # 控制处理速率
                    if self.rate_check:
                        time.sleep(1 / self.rate)

                    # 初始化数据源时，强制指定摄像头后端
                    if self.source.isnumeric():
                        cap = cv2.VideoCapture(int(self.source), cv2.CAP_MSMF)  # 使用MSMF后端
                    else:
                        cap = cv2.VideoCapture(self.source)

        except Exception as e:
            self.send_msg.emit(f'错误: {str(e)}')



class MainWindow(QMainWindow, Ui_mainWindow):
    def __init__(self, parent=None):
        super(MainWindow, self).__init__(parent)
        self.setupUi(self)
        self.m_flag = False
        # 无边框
        self.setWindowFlags(Qt.FramelessWindowHint)
        # ---gtj 设置窗口透明
        self.setAttribute(Qt.WA_TranslucentBackground)
        # 新增：禁用QLabel自动拉伸
        self.out_video.setScaledContents(False)
        self.raw_video.setScaledContents(False)
        # 窗口控制
        self.minButton.clicked.connect(self.showMinimized)
        self.maxButton.clicked.connect(self.max_or_restore)
        self.closeButton.clicked.connect(self.close)
        # 模型列表初始化
        self.pt_list = [f for f in os.listdir('./pt') if f.endswith('.pt')]
        self.pt_list.sort(key=lambda x: os.path.getsize('./pt/' + x))
        self.modelBox.clear()  # 修改为win4的控件名
        self.modelBox.addItems(self.pt_list)  # 修改为win4的控件名
        # 检测线程连接
        self.det_thread = DetThread()
        self.det_thread.send_img.connect(lambda x: self.show_image(x, self.out_video))
        self.det_thread.send_raw.connect(lambda x: self.show_image(x, self.raw_video))
        self.det_thread.send_statistic.connect(self.show_statistic)
        self.det_thread.send_msg.connect(self.show_msg)
        self.det_thread.send_percent.connect(self.progressBar.setValue)
        self.det_thread.send_fps.connect(self.fps_label.setText)
        # 按钮连接
        self.fileButton.clicked.connect(self.open_file)
        self.cameraButton.clicked.connect(self.chose_cam)
        self.runButton.clicked.connect(self.run_or_continue)
        self.stopButton.clicked.connect(self.stop)
        self.modelBox.currentTextChanged.connect(self.change_model)  # 修改为win4的控件名
        # 参数控制
        self.weedSpinBox.valueChanged.connect(lambda x: self.change_val(x, 'weedSpinBox'))
        self.weedSlider.valueChanged.connect(lambda x: self.change_val(x, 'weedSlider'))
        self.confSpinBox.valueChanged.connect(lambda x: self.change_val(x, 'confSpinBox'))
        self.confSlider.valueChanged.connect(lambda x: self.change_val(x, 'confSlider'))
        self.iouSpinBox.valueChanged.connect(lambda x: self.change_val(x, 'iouSpinBox'))
        self.iouSlider.valueChanged.connect(lambda x: self.change_val(x, 'iouSlider'))
        self.timeSpinBox.valueChanged.connect(lambda x: self.change_val(x, 'timeSpinBox'))  # 修改为win4的控件名
        self.timeSlider.valueChanged.connect(lambda x: self.change_val(x, 'timeSlider'))  # 修改为win4的控件名
        self.checkBox.clicked.connect(self.checkrate)

        self.load_setting()

        self.reportpushButton.clicked.connect(self.generate_report)
        # 初始化报告数据存储
        self.report_data = {
            "timestamp": "",
            "statistics": {},
            "weed_coords": [],
            "weed_density": ""
        }


    def reset_det_thread(self):
        # 保存旧线程的配置参数
        old_source = self.det_thread.source if hasattr(self.det_thread, 'source') else '0'
        old_weights = self.det_thread.weights if hasattr(self.det_thread, 'weights') else './yolo11s.pt'
        old_conf = self.det_thread.conf_thres if hasattr(self.det_thread, 'conf_thres') else 0.45
        old_iou = self.det_thread.iou_thres if hasattr(self.det_thread, 'iou_thres') else 0.70
        old_rate = self.det_thread.rate if hasattr(self.det_thread, 'rate') else 100
        old_weed_threshold = self.weedSpinBox.value()  # 保存阈值

        # 断开旧线程信号
        try:
            self.det_thread.send_img.disconnect()
            self.det_thread.send_raw.disconnect()
            self.det_thread.send_statistic.disconnect()
            self.det_thread.send_msg.disconnect()
            self.det_thread.send_percent.disconnect()
            self.det_thread.send_fps.disconnect()
        except TypeError:
            pass

        # 创建新线程并继承配置
        self.det_thread = DetThread()
        self.det_thread.source = old_source  # 继承数据源
        self.det_thread.weights = old_weights  # 继承模型权重
        self.det_thread.conf_thres = old_conf  # 继承置信度阈值
        self.det_thread.iou_thres = old_iou  # 继承IoU阈值
        self.det_thread.rate = old_rate  # 继承处理速率
        self.weedSpinBox.setValue(old_weed_threshold)  # 恢复阈值

        # 重新连接信号
        self.det_thread.send_img.connect(lambda x: self.show_image(x, self.out_video))
        self.det_thread.send_raw.connect(lambda x: self.show_image(x, self.raw_video))
        self.det_thread.send_statistic.connect(self.show_statistic)
        self.det_thread.send_msg.connect(self.show_msg)
        self.det_thread.send_percent.connect(self.progressBar.setValue)
        self.det_thread.send_fps.connect(self.fps_label.setText)

    def max_or_restore(self):
        if self.maxButton.isChecked():
            self.showMaximized()
        else:
            self.showNormal()

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.m_flag = True
            self.m_Position = event.pos()

    def mouseMoveEvent(self, event):
        if Qt.LeftButton and self.m_flag:
            self.move(event.globalPos() - self.m_Position)

    def mouseReleaseEvent(self, event):
        self.m_flag = False

    def search_pt(self):
        pt_list = [f for f in os.listdir('./pt') if f.endswith('.pt')]
        if pt_list != self.pt_list:
            self.pt_list = pt_list
            self.modelBox.clear()  # 修改为win4的控件名
            self.modelBox.addItems(pt_list)  # 修改为win4的控件名

    def checkrate(self):
        self.det_thread.rate_check = self.checkBox.isChecked()


    def chose_cam(self):
        try:
            self.stop()
            MessageBox(self.closeButton, title='提示', text='检测摄像头中...', time=2000, auto=True).exec_()
            _, cams = Camera().get_cam_num()
            menu = QMenu()
            actions = [QAction(cam, self) for cam in cams]
            for action in actions:
                menu.addAction(action)
            pos = self.cameraButton.mapToGlobal(QPoint(0, self.cameraButton.height()))
            action = menu.exec_(pos)
            if action:
                # 关键修改：保持source为字符串类型（如"0"）
                self.det_thread.source = action.text()  # 直接使用字符串（如"0"）
                self.statistic_msg(f'摄像头: {action.text()}')
        except Exception as e:
            self.statistic_msg(str(e))

    def load_setting(self):
        config_file = 'config/setting.json'
        default_config = {"iou": 0.70, "conf": 0.45, "rate": 10, "check": 0}
        if os.path.exists(config_file):
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
        else:
            config = default_config
            with open(config_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=2)
        self.confSpinBox.setValue(config['conf'])  # confSpinBox对应conf参数
        self.iouSpinBox.setValue(config['iou'])  # iouSpinBox对应iou参数
        self.timeSpinBox.setValue(config['rate'])  # 修改为win4的控件名
        self.checkBox.setChecked(bool(config['check']))
        self.det_thread.rate_check = bool(config['check'])

    def change_val(self, x, flag):
        if flag == 'confSpinBox':
            self.confSlider.setValue(int(x * 100))
            self.det_thread.conf_thres = x
        elif flag == 'confSlider':
            self.confSpinBox.setValue(x / 100)
        elif flag == 'iouSpinBox':
            self.iouSlider.setValue(int(x * 100))
            self.det_thread.iou_thres = x
        elif flag == 'iouSlider':
            self.iouSpinBox.setValue(x / 100)
        elif flag == 'timeSpinBox':  # 修改为win4的控件名
            self.timeSlider.setValue(x)  # 修改为win4的控件名
            self.det_thread.rate = x * 10
        elif flag == 'timeSlider':  # 修改为win4的控件名
            self.timeSpinBox.setValue(x)  # 修改为win4的控件名
        elif flag == 'weedSpinBox':
            self.weedSlider.setValue(int(x))  # 同步滑块位置
        elif flag == 'weedSlider':
            self.weedSpinBox.setValue(x)  # 同步数值框

    def statistic_msg(self, msg):
        self.downLabel.setText(msg)  # 修改为win4的控件名

    def show_msg(self, msg):
        self.runButton.setChecked(False)
        self.statistic_msg(msg)

    def change_model(self, x):
        self.det_thread.weights = f"./pt/{x}"
        self.statistic_msg(f'模型切换: {x}')

    def open_file(self):
        config_file = 'config/fold.json'
        default_config = {"open_fold": "./"}
        if os.path.exists(config_file):
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
        else:
            config = default_config
            with open(config_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=2)
        path, _ = QFileDialog.getOpenFileName(
            self, '选择文件', config['open_fold'],
            "媒体文件 (*.mp4 *.mkv *.avi *.flv *.jpg *.png)")
        if path:
            self.det_thread.source = path
            self.statistic_msg(f'文件: {os.path.basename(path)}')
            config['open_fold'] = os.path.dirname(path)
            with open(config_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=2)
            self.stop()

    def run_or_continue(self):
        # 如果线程已结束，先重置线程（继承参数）
        if not self.det_thread.isRunning():
            self.reset_det_thread()  # 重置但保留配置

        self.det_thread.jump_out = False
        if self.runButton.isChecked():
            self.det_thread.start()
            source = os.path.basename(self.det_thread.source)
            source = '摄像头' if source.isnumeric() else source
            self.statistic_msg(f'检测中: {source}')
        else:
            self.det_thread.is_continue = False
            self.statistic_msg('暂停中')

    def stop(self):
        self.det_thread.jump_out = True  # 终止线程循环
        self.det_thread.quit()  # 确保线程退出
        self.det_thread.wait()  # 等待线程结束
        self.statistic_msg('已停止')

    @staticmethod
    def show_image(img_src, label):
        try:
            #img_rgb = cv2.cvtColor(img_src, cv2.COLOR_BGR2RGB)
            #h, w, ch = img_rgb.shape
            h, w, ch = img_src.shape
            bytes_per_line = ch * w
            #q_img = QImage(img_rgb.data, w, h, bytes_per_line, QImage.Format_RGB888)
            q_img = QImage(img_src.data, w, h, bytes_per_line, QImage.Format_BGR888)
            pixmap = QPixmap.fromImage(q_img)
            # 缩放并保持宽高比，适应Label大小
            pixmap = pixmap.scaled(label.size(), Qt.KeepAspectRatio, Qt.SmoothTransformation)
            label.setPixmap(pixmap)
            label.setAlignment(Qt.AlignCenter)  # 图像居中显示
        except Exception as e:
            print(f"显示错误: {str(e)}")

    def show_statistic(self, statistic_dic, weed_coords):
        self.resultWidget.clear()
        # 更新报告数据
        self.report_data["statistics"] = statistic_dic
        self.report_data["weed_coords"] = weed_coords
        self.report_data["timestamp"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # 定义杂草类别
        weed_classes = ["broadleaf_weed", "grass_weed", "sedge_weed"]

        # 计算总杂草数量
        total = sum(statistic_dic.values())
        weed_count = sum(statistic_dic.get(cls, 0) for cls in weed_classes)
        weed_percent = (weed_count / total * 100) if total > 0 else 0.0

        # 密度描述判断
        if weed_percent >= 40:
            density_desc = "过多"
        elif weed_percent >= 20:
            density_desc = "较多"
        elif weed_percent >= 10:
            density_desc = "适中"
        else:
            density_desc = "较少"
        # 生成统计结果
        results = [f"{k}: {v}" for k, v in sorted(statistic_dic.items(),
                                                  key=lambda x: x[1], reverse=True) if v > 0]
        # 添加密度行
        results.append(f"总杂草密度：{weed_percent:.1f}% ({density_desc})")
        # 添加警报信息（如果超过阈值）
        alert_threshold = self.weedSpinBox.value()
        if weed_percent >= alert_threshold:
            results.append(f"超过杂草密度警报阈值({alert_threshold}%)，急需除草！")
        # 添加坐标信息
        if weed_coords:
            results.append("杂草坐标：")
            for coord in weed_coords:
                x, y = coord
                results.append(f"({x:.1f}, {y:.1f})")
        else:
            results.append("杂草坐标：无")

        self.resultWidget.addItems(results)
        # 提取杂草密度信息
        for item in results:
            if "总杂草密度" in item:
                self.report_data["weed_density"] = item
                break

    def closeEvent(self, event):
        self.det_thread.jump_out = True
        config = {
            "iou": self.iouSpinBox.value(),    # 正确关联iou
            "conf": self.confSpinBox.value(),  # 正确关联conf
            "rate": self.timeSpinBox.value(),
            "check": int(self.checkBox.isChecked())
        }
        with open('config/setting.json', 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=2)
        MessageBox(self.closeButton, title='提示', text='正在关闭...', time=2000, auto=True).exec_()
        event.accept()

    def generate_report(self):
        """生成统计报告"""
        try:
            # 创建报告目录
            report_dir = "./report"
            os.makedirs(report_dir, exist_ok=True)
            # 生成时间戳作为文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            txt_path = os.path.join(report_dir, f"weed_report_{timestamp}.txt")
            pdf_path = os.path.join(report_dir, f"weed_report_{timestamp}.pdf")
            word_path = os.path.join(report_dir, f"weed_report_{timestamp}.docx")  # 新增Word路径
            # 生成TXT报告
            self.generate_txt_report(txt_path)
            # 生成PDF报告
            self.generate_pdf_report(pdf_path)
            # 生成Word报告
            self.generate_word_report(word_path)  # 新增Word报告生成
            # 显示成功消息
            self.statistic_msg(f"报告已生成: {txt_path}, {pdf_path} 和 {word_path}")
        except Exception as e:
            self.statistic_msg(f"生成报告失败: {str(e)}")

    def generate_txt_report(self, file_path):
        """生成文本格式的报告"""
        with open(file_path, 'w', encoding='utf-8') as f:
            # 报告标题
            f.write("=" * 60 + "\n")
            f.write(f"杂草识别与定位统计报告\n".center(60))
            f.write("=" * 60 + "\n\n")
            # 报告时间
            f.write(f"生成时间: {self.report_data['timestamp']}\n\n")
            # 统计信息
            f.write("=" * 60 + "\n")
            f.write("植物种类统计:\n")
            f.write("=" * 60 + "\n")
            for plant, count in self.report_data['statistics'].items():
                f.write(f"{plant:<20}: {count:>4}株\n")
            # 杂草密度
            f.write("\n" + "=" * 60 + "\n")
            f.write("杂草密度分析:\n")
            f.write("=" * 60 + "\n")
            f.write(f"{self.report_data['weed_density']}\n\n")
            # 坐标信息
            f.write("=" * 60 + "\n")
            f.write("杂草中心点坐标:\n")
            f.write("=" * 60 + "\n")
            if self.report_data['weed_coords']:
                for i, (x, y) in enumerate(self.report_data['weed_coords'], 1):
                    f.write(f"杂草 {i:03}: X={x:.1f}, Y={y:.1f}\n")
            else:
                f.write("未检测到杂草\n")
            # 报告结尾
            f.write("\n" + "=" * 60 + "\n")
            f.write("报告结束\n".center(60))
            f.write("=" * 60 + "\n")

    def generate_pdf_report(self, file_path):
        """生成PDF格式的报告（支持中文）"""
        try:
            # 创建fonts目录存放字体
            font_dir = "./fonts"
            os.makedirs(font_dir, exist_ok=True)
            # 优先尝试使用程序自带字体
            font_path = None
            for font_name in ["simhei.ttf", "simsun.ttc", "msyh.ttf"]:
                local_font = os.path.join(font_dir, font_name)
                if os.path.exists(local_font):
                    font_path = local_font
                    break
            # 如果程序目录没有字体，尝试使用系统字体
            if not font_path:
                # Windows系统字体路径
                win_fonts = [
                    "C:/Windows/Fonts/simhei.ttf",
                    "C:/Windows/Fonts/simsun.ttc",
                    "C:/Windows/Fonts/msyh.ttf"
                ]
                # Linux系统字体路径
                linux_fonts = [
                    "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
                    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
                    "/usr/share/fonts/truetype/arphic/uming.ttc"
                ]
                # macOS系统字体路径
                mac_fonts = [
                    "/Library/Fonts/Arial Unicode.ttf",
                    "/System/Library/Fonts/PingFang.ttc"
                ]
                # 尝试所有可能的字体路径
                possible_fonts = win_fonts + linux_fonts + mac_fonts
                for path in possible_fonts:
                    if os.path.exists(path):
                        font_path = path
                        break
            # 如果仍然找不到字体，使用默认字体（中文可能显示为方框）
            if not font_path:
                return self._generate_pdf_with_default_font(file_path)
            # 获取字体名称（不带扩展名）
            font_name = os.path.basename(font_path).split('.')[0]
            # 2. 创建PDF对象
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            # 3. 添加中文字体
            pdf.add_font(font_name, '', font_path, uni=True)
            pdf.set_font(font_name, size=12)
            # 4. 生成中文报告内容
            # 报告标题
            pdf.set_font_size(16)
            pdf.cell(0, 10, "杂草识别与定位统计报告", 0, 1, 'C')
            pdf.ln(5)
            # 报告时间
            pdf.set_font_size(12)
            pdf.cell(0, 10, f"生成时间: {self.report_data['timestamp']}", 0, 1)
            pdf.ln(5)
            # 统计信息
            pdf.set_font_size(14)
            pdf.cell(0, 10, "植物种类统计:", 0, 1)
            pdf.set_font_size(12)
            pdf.set_fill_color(240, 240, 240)
            pdf.cell(60, 10, "植物种类", 1, 0, 'C', True)
            pdf.cell(60, 10, "数量", 1, 1, 'C', True)
            for plant, count in self.report_data['statistics'].items():
                pdf.cell(60, 10, plant, 1)
                pdf.cell(60, 10, str(count), 1, 1)
            pdf.ln(5)
            # 杂草密度
            pdf.set_font_size(14)
            pdf.cell(0, 10, "杂草密度分析:", 0, 1)
            pdf.set_font_size(12)
            pdf.multi_cell(0, 10, self.report_data['weed_density'])
            pdf.ln(5)
            # 坐标信息
            pdf.set_font_size(14)
            pdf.cell(0, 10, "杂草中心点坐标:", 0, 1)
            pdf.set_font_size(12)
            if self.report_data['weed_coords']:
                pdf.set_fill_color(240, 240, 240)
                pdf.cell(40, 10, "杂草编号", 1, 0, 'C', True)
                pdf.cell(60, 10, "X坐标", 1, 0, 'C', True)
                pdf.cell(60, 10, "Y坐标", 1, 1, 'C', True)

                for i, (x, y) in enumerate(self.report_data['weed_coords'], 1):
                    pdf.cell(40, 10, f"杂草 {i:03}", 1)
                    pdf.cell(60, 10, f"{x:.1f}", 1)
                    pdf.cell(60, 10, f"{y:.1f}", 1, 1)
            else:
                pdf.cell(0, 10, "未检测到杂草", 0, 1)
            # 保存PDF
            pdf.output(file_path)
            return True
        except Exception as e:
            error_msg = f"PDF生成错误: {str(e)}\n{traceback.format_exc()}"
            self.statistic_msg(error_msg)
            print(error_msg)
            return False

    def generate_word_report(self, file_path):
        """生成Word格式的报告（.docx）"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.oxml.ns import qn
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.enum.table import WD_TABLE_ALIGNMENT
            # 创建文档
            doc = Document()
            # 设置中文字体
            doc.styles['Normal'].font.name = u'宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            # 标题
            title = doc.add_heading('杂草识别与定位统计报告', level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # 报告时间
            time_para = doc.add_paragraph()
            time_para.add_run(f"生成时间: {self.report_data['timestamp']}").bold = True
            doc.add_paragraph()
            # 植物种类统计
            doc.add_heading('植物种类统计', level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Shading Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '植物种类'
            hdr_cells[1].text = '数量'

            for plant, count in self.report_data['statistics'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = plant
                row_cells[1].text = str(count)

            doc.add_paragraph()
            # 杂草密度
            doc.add_heading('杂草密度分析', level=1)
            density_para = doc.add_paragraph()
            density_para.add_run(self.report_data["weed_density"]).bold = True
            doc.add_paragraph()
            # 杂草坐标
            doc.add_heading('杂草中心点坐标', level=1)
            if self.report_data['weed_coords']:
                coord_table = doc.add_table(rows=1, cols=3)
                coord_table.style = 'Light Shading Accent 1'
                coord_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr_cells = coord_table.rows[0].cells
                hdr_cells[0].text = '杂草编号'
                hdr_cells[1].text = 'X坐标'
                hdr_cells[2].text = 'Y坐标'

                for i, (x, y) in enumerate(self.report_data['weed_coords'], 1):
                    row_cells = coord_table.add_row().cells
                    row_cells[0].text = f"杂草 {i:03}"
                    row_cells[1].text = f"{x:.1f}"
                    row_cells[2].text = f"{y:.1f}"
            else:
                doc.add_paragraph('未检测到杂草')
            # 添加系统信息
            doc.add_paragraph()
            doc.add_paragraph("报告生成系统: 杂草识别、定位及状况监测系统")
            doc.add_paragraph(f"模型版本: {os.path.basename(self.det_thread.weights)}")
            # 保存文档
            doc.save(file_path)
            return True
        except ImportError:
            self.statistic_msg("生成Word报告失败: 请安装python-docx库")
            return False
        except Exception as e:
            self.statistic_msg(f"生成Word报告失败: {str(e)}")
            return False


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())